import streamlit as st
from google import genai
from google.genai import types
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import io
import time

# ==========================================
# 1. 워드 파일 디자인 & 생성 헬퍼 함수
# ==========================================
def set_cell_borders(cell, color="D9D9D9", sz="4", val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
    tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def create_word_document(all_word_data):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("통합 본문 단어장")
    title_run.font.bold = True
    title_run.font.size = Pt(18)
    title_p.paragraph_format.space_after = Pt(24)
    
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    col_widths = [Inches(1.5), Inches(2.0), Inches(4.0)]
    
    hdr_cells = table.rows[0].cells
    headers = ["본문 단어", "우리말 뜻", "영영 풀이"]
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        hdr_cells[i].width = col_widths[i]
        set_cell_shading(hdr_cells[i], "4F81BD")
        set_cell_borders(hdr_cells[i], color="A6A6A6")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = docx.shared.RGBColor(255, 255, 255)
            
    for row_idx, item in enumerate(all_word_data):
        row_cells = table.add_row().cells
        row_data = [item.get("word", ""), item.get("meaning", ""), item.get("definition", "")]
        
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            row_cells[i].width = col_widths[i]
            if row_idx % 2 == 1:
                set_cell_shading(row_cells[i], "F2F5F8")
            set_cell_borders(row_cells[i], color="D9D9D9")
            
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(10)
                
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def process_images_safely(client, uploaded_files, api_key, progress_bar, status_text):
    all_data = []
    total_files = len(uploaded_files)
    
    prompt = """
    이 이미지에서 영어 단어, 우리말 뜻, 영영 풀이를 추출해서 정확한 JSON 배열 형식으로 출력해줘.
    필기구로 수정한 흔적이나 추가로 적은 필기는 무시하고, 원래 인쇄되어 있던 텍스트만 추출해줘.
    결과는 오직 아래 구조를 가진 JSON 데이터만 반환해야 해:
    [
      {"word": "단어", "meaning": "품사 및 뜻", "definition": "영영 풀이 내용"}
    ]
    """
    
    current_displayed_percent = 0
    
    for idx, file in enumerate(uploaded_files):
        target_percent = int(((idx + 1) / total_files) * 100)
        
        pre_target = target_percent - 3 if target_percent > 3 else 0
        while current_displayed_percent < pre_target:
            current_displayed_percent += 1
            progress_bar.progress(current_displayed_percent / 100)
            status_text.markdown(f"<div class='status-msg'>🌱 단어 아카이브 분석 중... {current_displayed_percent}% (`{file.name}` 판독 중)</div>", unsafe_allow_html=True)
            time.sleep(0.01)
            
        page_data = None
        max_retries = 3
        
        for attempt in range(max_retries):
            try:
                file.seek(0)
                image_bytes = file.read()
                
                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=[
                        types.Part.from_bytes(data=image_bytes, mime_type=file.type),
                        prompt
                    ],
                    config=types.GenerateContentConfig(response_mime_type="application/json")
                )
                page_data = json.loads(response.text)
                all_data.extend(page_data)
                break
                
            except Exception as e:
                if attempt < max_retries - 1:
                    status_text.markdown(f"<div class='status-msg' style='color:#cca01a;'>⚠️ 안정한 연결을 위해 재시도 중입니다... ({attempt + 1}/{max_retries})</div>", unsafe_allow_html=True)
                    time.sleep(2.0)
                else:
                    st.error(f"🛑 구글 서버 과부하가 지속되어 `{file.name}` 처리에 실패했습니다. 잠시 후 다시 시도해 주세요.")
            
        while current_displayed_percent < target_percent:
            current_displayed_percent += 1
            progress_bar.progress(current_displayed_percent / 100)
            status_text.markdown(f"<div class='status-msg'>✨ 단어 아카이브 분석 중... {current_displayed_percent}% (`{file.name}` 정제 완료)</div>", unsafe_allow_html=True)
            time.sleep(0.01)
            
        time.sleep(0.2)
        
    progress_bar.progress(1.0)
    status_text.markdown("<div class='status-msg' style='color:#2e7d32; font-weight:bold;'>🌿 정제 프로세스가 완료되었습니다. (100%)</div>", unsafe_allow_html=True)
    return all_data

# ==========================================
# 3. Streamlit 메인 UI 대시보드
# ==========================================
st.set_page_config(page_title="Voca-converter", layout="wide", page_icon="📝")

# 🔥 [치트키] 캐시를 무시하고 웹 화면 전체를 강제로 지배하는 초고도 CSS 주입
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&family=Noto+Sans+KR:wght@300;400;500;700&display=swap');
    
    /* 1. 텅 빈 좌우 여백을 채우기 위해 기본 앱 래퍼 크기 강제 확장 */
    [data-testid="stAppViewContainer"] {
        background-color: #ECEFF1 !important; /* 부드러운 오피스 모노톤 그레이 */
        font-family: 'Noto Sans KR', 'Inter', sans-serif !important;
    }
    
    /* 2. 상단 지저분한 여백과 기본 흰색 바 강제 삭제 및 중앙 카드화 */
    [data-testid="stMainBlockContainer"] {
        background-color: #FFFFFF !important;
        max-width: 850px !important;
        margin: 60px auto !important;
        padding: 60px 60px !important;
        border-radius: 24px !important;
        box-shadow: 0 15px 35px rgba(0, 0, 0, 0.05) !important;
    }
    
    /* 3. 상단 헤더 박스 라인 처리 (학원 교재 연구실 무드) */
    .header-box {
        text-align: center !important;
        border-bottom: 2px solid #F1F5F9 !important;
        padding-bottom: 30px !important;
        margin-bottom: 40px !important;
    }
    .main-title {
        font-family: 'Inter', sans-serif !important;
        font-size: 40px !important;
        font-weight: 600 !important;
        color: #0F172A !important;
        letter-spacing: -1.5px !important;
    }
    .sub-title {
        font-family: 'Inter', sans-serif !important;
        font-size: 11px !important;
        color: #94A3B8 !important;
        font-weight: 600 !important;
        letter-spacing: 3px !important;
        margin-top: 8px !important;
    }
    
    /* 4. 세련된 라운드 설명 블록 */
    .description-text {
        font-size: 14.5px !important;
        color: #475569 !important;
        line-height: 1.7 !important;
        margin-bottom: 40px !important;
        background-color: #F8FAFC !important;
        padding: 22px 30px !important;
        border-radius: 14px !important;
        border: 1px solid #E2E8F0 !important;
        text-align: center !important;
    }
    
    /* 5. 투박한 업로드 점선 상자를 완전히 심플하고 감성적인 영역으로 재구축 */
    [data-testid="stFileUploader"] {
        border: 2px dashed #CBD5E1 !important;
        background-color: #FAFAFA !important;
        border-radius: 16px !important;
        padding: 30px 20px !important;
    }
    
    /* 6. 투박한 기본 버튼 스타일을 지우고 30대 원장 취향의 매트 블랙으로 고도화 */
    .stButton>button {
        background-color: #1E293B !important;
        color: white !important;
        border-radius: 10px !important;
        padding: 15px 20px !important;
        font-size: 16px !important;
        font-weight: 500 !important;
        border: none !important;
        box-shadow: 0 4px 12px rgba(30, 41, 59, 0.15) !important;
        width: 100% !important;
        margin-top: 25px !important;
    }
    .stButton>button:hover {
        background-color: #0F172A !important;
        box-shadow: 0 6px 20px rgba(15, 23, 42, 0.2) !important;
    }
    
    /* 7. 청록색 감성의 다운로드 버튼 */
    [data-testid="stDownloadButton"]>button {
        background-color: #0D9488 !important;
        color: white !important;
        border-radius: 10px !important;
        padding: 15px 20px !important;
        font-size: 16px !important;
        font-weight: 500 !important;
        border: none !important;
    }
    
    /* 진행률 바 */
    .status-msg {
        font-size: 14px !important;
        color: #475569 !important;
        margin-top: 25px !important;
    }
    .stProgress > div > div > div > div {
        background-color: #475569 !important;
    }
    </style>
""", unsafe_allow_html=True)

# 🏷️ 상단 감성 헤더 디자인 컴포넌트
st.markdown("""
    <div class='header-box'>
        <div class='main-title'>Voca-converter</div>
        <div class='sub-title'>MADE BY MANJU</div>
    </div>
""", unsafe_allow_html=True)

st.markdown("<div class='description-text'>교재나 유인물 사진을 업로드하시면, 수업에 즉시 활용할 수 있는 단정하고 정돈된 <strong>표 형태의 워드 문서(.docx)</strong>로 통합 변환해 드립니다.</div>", unsafe_allow_html=True)

if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    st.error("❌ Streamlit Cloud 설정의 Secrets에 GEMINI_API_KEY가 등록되지 않았습니다.")
    st.stop()

uploaded_files = st.file_uploader(
    "교재 및 단어장 사진 파일을 선택하세요 (여러 장 동시 선택 가능)", 
    type=["jpg", "jpeg", "png"], 
    accept_multiple_files=True
)

if uploaded_files:
    st.markdown(f"<div style='font-size:13.5px; color:#64748B; margin-bottom:15px; text-align:center;'>📂 <strong>선택된 아카이브:</strong> 총 {len(uploaded_files)}개의 문서 파일이 대기 중입니다.</div>", unsafe_allow_html=True)
    
    if st.button("✨ 업로드된 문서 분석 및 Word 파일 생성"):
        client = genai.Client(api_key=api_key)
        
        status_text = st.empty()
        progress_bar = st.progress(0)
        
        all_word_data = process_images_safely(client, uploaded_files, api_key, progress_bar, status_text)
        
        if all_word_data:
            st.success("🎉 모든 단어 데이터 정제가 성공적으로 완료되었습니다!")
            st.write("### 🔍 데이터 통합 미리보기")
            st.dataframe(all_word_data, use_container_width=True)
            
            word_file_buffer = create_word_document(all_word_data)
            
            st.markdown("<div style='margin-top: 25px;'></div>", unsafe_allow_html=True)
            st.download_button(
                label="📥 정제된 수업용 Word 문서 다운로드 (.docx)",
                data=word_file_buffer,
                file_name="🔮_통합_영어단어장.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
